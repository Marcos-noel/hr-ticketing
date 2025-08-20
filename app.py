from flask import Flask, render_template, request, jsonify, redirect, url_for, session, send_file
import mysql.connector
from datetime import datetime, timedelta
from mysql.connector import Error
import smtplib
from email.mime.text import MIMEText
import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import openpyxl
from io import BytesIO
import math

app = Flask(__name__)
app.config['SECRET_KEY'] = 'keproba-secret-key'
app.config['SMTP_SERVER'] = 'smtp.gmail.com'
app.config['SMTP_PORT'] = 587
app.config['SMTP_USER'] = 'ICT2@brand.ke'
app.config['RECIPIENT_EMAILS'] = ['SWanjiku@brand.ke', 'NWathika@brand.ke', 'Slosuko@brand.ke']
app.config['SMTP_PASSWORD'] = 'your-app-password' # REPLACE THIS

def create_db_connection():
    """Create a MySQL database connection."""
    connection = None
    try:
        connection = mysql.connector.connect(
            host='localhost',
            port=3307,
            user='root',
            password='3306',
            database='HR_helpdesk'
        )
        if connection.is_connected():
            return connection
    except Error as e:
        print(f"Error connecting to MySQL: {e}")
    return connection

def init_database():
    """Create the database and tables if they don't exist, and insert sample data."""
    temp_connection = None
    temp_cursor = None
    connection = None
    cursor = None
    try:
        temp_connection = mysql.connector.connect(
            host='localhost',
            port=3307,
            user='root',
            password='3306'
        )
        temp_cursor = temp_connection.cursor()
        temp_cursor.execute("CREATE DATABASE IF NOT EXISTS HR_helpdesk")
        print("Database 'HR_helpdesk' created or already exists.")
        
        connection = mysql.connector.connect(
            host='localhost',
            port=3307,
            user='root',
            password='3306',
            database='HR_helpdesk'
        )
        cursor = connection.cursor()
        print("Database 'HR_helpdesk' ready.")

        tables = {}
        tables['tickets'] = """
        CREATE TABLE IF NOT EXISTS tickets (
            id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(100) NOT NULL,
            department VARCHAR(100) NOT NULL,
            issue_type VARCHAR(100) NOT NULL,
            priority VARCHAR(50) NOT NULL,
            description TEXT NOT NULL,
            status VARCHAR(50) DEFAULT 'Open',
            assigned_to VARCHAR(100),
            resolution_note TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            sla_deadline DATETIME,
            resolution_note_date DATETIME
        )
        """
        tables['ticket_history'] = """
        CREATE TABLE IF NOT EXISTS ticket_history (
            id INT AUTO_INCREMENT PRIMARY KEY,
            ticket_id INT,
            action VARCHAR(100) NOT NULL,
            performed_by VARCHAR(100) NOT NULL,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (ticket_id) REFERENCES tickets(id) ON DELETE CASCADE
        )
        """
        tables['hr_staff'] = """
        CREATE TABLE IF NOT EXISTS hr_staff (
            id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(100) NOT NULL,
            email VARCHAR(100) NOT NULL,
            role VARCHAR(50) DEFAULT 'Staff'
        )
        """
        
        for table_name, table_sql in tables.items():
            cursor.execute(table_sql)
            print(f"Table '{table_name}' created or already exists.")

        cursor.execute("SELECT COUNT(*) FROM hr_staff")
        if cursor.fetchone()[0] == 0:
            staff = [
                ('S. Wanjiku', 'SWanjiku@brand.ke', 'Manager'),
                ('N. Wathika', 'NWathika@brand.ke', 'Staff'),
                ('S. Losuko', 'Slosuko@brand.ke', 'Staff')
            ]
            for s in staff:
                cursor.execute("INSERT INTO hr_staff (name, email, role) VALUES (%s, %s, %s)", s)
            connection.commit()
            print("Database initialized successfully with sample HR staff")
        
    except Error as e:
        print(f"Error initializing database: {e}")
    finally:
        if temp_cursor:
            temp_cursor.close()
        if temp_connection and temp_connection.is_connected():
            temp_connection.close()
        if cursor:
            cursor.close()
        if connection and connection.is_connected():
            connection.close()

def send_email(to_emails, subject, body):
    try:
        msg = MIMEText(body)
        msg['Subject'] = subject
        msg['From'] = app.config['SMTP_USER']
        
        with smtplib.SMTP(app.config['SMTP_SERVER'], app.config['SMTP_PORT']) as server:
            server.starttls()
            server.login(app.config['SMTP_USER'], app.config['SMTP_PASSWORD'])
            server.sendmail(app.config['SMTP_USER'], to_emails, msg.as_string())
        return True
    except Exception as e:
        print(f"Email error: {e}")
        return False

def get_current_user():
    return session.get('user', 'HR Staff')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username', '')
        if username:
            session['user'] = username
            send_email(app.config['RECIPIENT_EMAILS'], "KEPROBA HR Login", f"Welcome, {username}!")
            return redirect(url_for('hr_dashboard'))
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('user', None)
    return redirect(url_for('login'))

@app.route('/api/current_user')
def get_current_user_api():
    return jsonify({'user': get_current_user()})

@app.route('/')
def index():
    departments = [
        'Corporate Services', 'Research and Innovation', 'Product and Market Development',
        'Nation Branding', 'Resource Centre', 'Stakeholder Engagement',
        'Finance Department', 'ICT Department', 'HR Department'
    ]
    issue_types = [
        'training matters', 'leave management', 'payroll issue',
        'medical insurance issue','group life insurance', 'clearance matters',
        'mortgage scheme', 'car loan facility'
    ]
    priorities = ['Low', 'Medium', 'High']
    return render_template('index.html', departments=departments, issue_types=issue_types, priorities=priorities)

@app.route('/hr')
def hr_dashboard():
    if not session.get('user'):
        session['user'] = 'HR Staff'
    return render_template('hr.html')

@app.route('/api/tickets', methods=['GET', 'POST'])
def api_tickets():
    connection = create_db_connection()
    if not connection:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = connection.cursor(dictionary=True)
        if request.method == 'POST':
            data = request.get_json()
            if not data:
                return jsonify({'error': 'No data provided'}), 400
            if not data.get('name') or not data.get('description') or not data.get('priority'):
                return jsonify({'error': 'Name, Description, and Priority are required'}), 400
            
            department = data.get('department', 'Unassigned')
            issue_type = data.get('issue_type', 'Unassigned')
            if department == "Select Department...": department = "Unassigned"
            if issue_type == "Select Issue...": issue_type = "Unassigned"

            sla_hours = {'High': 24, 'Medium': 48, 'Low': 72}
            sla_deadline = datetime.utcnow() + timedelta(hours=sla_hours.get(data['priority'], 72))

            sql = """
            INSERT INTO tickets (name, department, issue_type, priority, description, sla_deadline)
            VALUES (%s, %s, %s, %s, %s, %s)
            """
            cursor.execute(sql, (data['name'], department, issue_type, data['priority'], data['description'], sla_deadline))
            new_ticket_id = cursor.lastrowid
            
            history_sql = "INSERT INTO ticket_history (ticket_id, action, performed_by, timestamp) VALUES (%s, %s, %s, %s)"
            cursor.execute(history_sql, (new_ticket_id, f"Ticket created by {data['name']}", data['name'], datetime.utcnow()))
            connection.commit()
            
            email_body = f"""
            New HR Ticket Submitted!

            Ticket ID: {new_ticket_id}
            Submitted By: {data['name']}
            Department: {department}
            Issue Type: {issue_type}
            Priority: {data['priority']}
            Description: {data['description']}
            SLA Deadline: {sla_deadline.strftime('%Y-%m-%d %H:%M:%S UTC')}

            View ticket: {url_for('hr_dashboard', _external=True)}
            """
            send_email(app.config['RECIPIENT_EMAILS'], f"New HR Ticket #{new_ticket_id}", email_body)
            return jsonify({'message': 'Ticket submitted successfully!', 'id': new_ticket_id}), 201

        elif request.method == 'GET':
            status = request.args.get('status', 'All')
            search = request.args.get('search', '')
            my_claims = request.args.get('my_claims', 'false').lower() == 'true'
            page = int(request.args.get('page', 1))
            page_size = int(request.args.get('page_size', 10))

            params = []
            where_clauses = []
            if status == 'Pending':
                where_clauses.append("status IN ('Open', 'In Progress')")
            elif status != 'All':
                where_clauses.append("status = %s")
                params.append(status)
            if search:
                where_clauses.append("(name LIKE %s OR department LIKE %s OR issue_type LIKE %s OR description LIKE %s)")
                search_param = f"%{search}%"
                params.extend([search_param, search_param, search_param, search_param])
            if my_claims:
                where_clauses.append("assigned_to = %s")
                params.append(get_current_user())

            where_sql = " WHERE " + " AND ".join(where_clauses) if where_clauses else ""
            sql = f"""
            SELECT t.*, 
                   GROUP_CONCAT(CONCAT(th.action, '||', th.performed_by, '||', DATE_FORMAT(th.timestamp, '%Y-%m-%d %H:%i')) 
                                ORDER BY th.timestamp SEPARATOR '##') as history_str
            FROM tickets t
            LEFT JOIN ticket_history th ON t.id = th.ticket_id
            {where_sql}
            GROUP BY t.id
            ORDER BY CASE WHEN priority = 'High' THEN 1 WHEN priority = 'Medium' THEN 2 ELSE 3 END, created_at DESC
            LIMIT %s OFFSET %s
            """
            params.extend([page_size, (page - 1) * page_size])
            cursor.execute(sql, params)
            tickets = cursor.fetchall()

            for ticket in tickets:
                if ticket['history_str']:
                    ticket['history'] = [
                        {'action': h.split('||')[0], 'performed_by': h.split('||')[1], 'timestamp': h.split('||')[2]}
                        for h in ticket['history_str'].split('##')
                    ]
                else:
                    ticket['history'] = []
                del ticket['history_str']
                ticket['created_at'] = ticket['created_at'].strftime('%Y-%m-%d %H:%M:%S') if ticket['created_at'] else None
                ticket['sla_deadline'] = ticket['sla_deadline'].strftime('%Y-%m-%d %H:%M:%S') if ticket['sla_deadline'] else None
                ticket['resolution_note_date'] = ticket['resolution_note_date'].strftime('%Y-%m-%d %H:%M:%S') if ticket['resolution_note_date'] else None

            count_sql = f"SELECT COUNT(*) as total FROM tickets {where_sql}"
            cursor.execute(count_sql, params[:-2])
            total = cursor.fetchone()['total']
            total_pages = math.ceil(total / page_size)

            return jsonify({
                'tickets': tickets,
                'total_pages': total_pages,
                'current_page': page
            })

    except Error as e:
        print(f"Database error: {e}")
        return jsonify({'error': 'Database error occurred'}), 500
    finally:
        if connection and connection.is_connected():
            connection.close()

@app.route('/api/tickets/<int:ticket_id>', methods=['PUT'])
def update_ticket(ticket_id):
    connection = create_db_connection()
    if not connection:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        data = request.get_json()
        current_user = get_current_user()
        cursor = connection.cursor()
        updates = []
        params = []
        history_action = None

        if 'assigned_to' in data:
            updates.append("assigned_to = %s")
            params.append(data['assigned_to'])
            history_action = f"Assigned to {data['assigned_to']}"
        if 'status' in data:
            updates.append("status = %s")
            params.append(data['status'])
            history_action = f"Status changed to {data['status']}"
            if data['status'] == 'Resolved':
                updates.append("resolution_note_date = %s")
                params.append(datetime.utcnow())
        if 'resolution_note' in data:
            updates.append("resolution_note = %s")
            params.append(data['resolution_note'])
            if history_action is None:
                history_action = "Resolution note updated"
        
        if not updates:
            return jsonify({'error': 'No valid fields to update'}), 400

        sql = f"UPDATE tickets SET {', '.join(updates)} WHERE id = %s"
        params.append(ticket_id)
        cursor.execute(sql, params)
        
        if history_action:
            history_sql = "INSERT INTO ticket_history (ticket_id, action, performed_by, timestamp) VALUES (%s, %s, %s, %s)"
            cursor.execute(history_sql, (ticket_id, history_action, current_user, datetime.utcnow()))
        connection.commit()

        email_body = f"Ticket #{ticket_id} has been updated by {current_user}.\n\nAction: {history_action}\n\nResolution Note: {data.get('resolution_note', 'N/A')}"
        send_email(app.config['RECIPIENT_EMAILS'], f"Update on HR Ticket #{ticket_id}", email_body)
        
        return jsonify({'message': 'Ticket updated successfully', 'ticket_id': ticket_id})

    except Error as e:
        print(f"Database error: {e}")
        return jsonify({'error': 'Database error occurred'}), 500
    finally:
        if connection and connection.is_connected():
            connection.close()

@app.route('/api/unclaimed_tickets')
def get_unclaimed_tickets():
    connection = create_db_connection()
    if not connection:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = connection.cursor(dictionary=True)
        sql = """
        SELECT t.*, GROUP_CONCAT(CONCAT(th.action, '||', th.performed_by, '||', DATE_FORMAT(th.timestamp, '%Y-%m-%d %H:%i')) ORDER BY th.timestamp SEPARATOR '##') as history_str
        FROM tickets t
        LEFT JOIN ticket_history th ON t.id = th.ticket_id
        WHERE t.assigned_to IS NULL AND t.status = 'Open'
        GROUP BY t.id
        ORDER BY t.priority DESC, t.created_at ASC
        """
        cursor.execute(sql)
        tickets = cursor.fetchall()
        for ticket in tickets:
            if ticket['history_str']:
                ticket['history'] = [
                    {'action': h.split('||')[0], 'performed_by': h.split('||')[1], 'timestamp': h.split('||')[2]}
                    for h in ticket['history_str'].split('##')
                ]
            else:
                ticket['history'] = []
            del ticket['history_str']
            ticket['created_at'] = ticket['created_at'].strftime('%Y-%m-%d %H:%M:%S') if ticket['created_at'] else None
            ticket['sla_deadline'] = ticket['sla_deadline'].strftime('%Y-%m-%d %H:%M:%S') if ticket['sla_deadline'] else None
            ticket['resolution_note_date'] = ticket['resolution_note_date'].strftime('%Y-%m-%d %H:%M:%S') if ticket['resolution_note_date'] else None
        return jsonify(tickets)
    except Error as e:
        print(f"Database error: {e}")
        return jsonify({'error': 'Database error occurred'}), 500
    finally:
        if connection and connection.is_connected():
            connection.close()

@app.route('/api/my_tickets', methods=['GET'])
def get_my_tickets():
    connection = create_db_connection()
    if not connection:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        name = request.args.get('name')
        if not name:
            return jsonify({'error': 'Name parameter is required'}), 400
        
        cursor = connection.cursor(dictionary=True)
        sql = """
        SELECT t.*, GROUP_CONCAT(CONCAT(th.action, '||', th.performed_by, '||', DATE_FORMAT(th.timestamp, '%Y-%m-%d %H:%i')) ORDER BY th.timestamp SEPARATOR '##') as history_str
        FROM tickets t
        LEFT JOIN ticket_history th ON t.id = th.ticket_id
        WHERE t.name = %s
        GROUP BY t.id
        ORDER BY t.created_at DESC
        """
        cursor.execute(sql, (name,))
        tickets = cursor.fetchall()
        
        for ticket in tickets:
            if ticket['history_str']:
                ticket['history'] = [
                    {'action': h.split('||')[0], 'performed_by': h.split('||')[1], 'timestamp': h.split('||')[2]}
                    for h in ticket['history_str'].split('##')
                ]
            else:
                ticket['history'] = []
            del ticket['history_str']
            ticket['created_at'] = ticket['created_at'].strftime('%Y-%m-%d %H:%M:%S') if ticket['created_at'] else None
            ticket['sla_deadline'] = ticket['sla_deadline'].strftime('%Y-%m-%d %H:%M:%S') if ticket['sla_deadline'] else None
            ticket['resolution_note_date'] = ticket['resolution_note_date'].strftime('%Y-%m-%d %H:%M:%S') if ticket['resolution_note_date'] else None
            
        return jsonify(tickets)
    except Error as e:
        print(f"Database error: {e}")
        return jsonify({'error': 'Database error occurred'}), 500
    finally:
        if connection and connection.is_connected():
            connection.close()

@app.route('/api/ticket_stats')
def get_ticket_stats():
    connection = create_db_connection()
    if not connection:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = connection.cursor(dictionary=True)
        
        stats = {}
        
        cursor.execute("SELECT COUNT(*) as total FROM tickets")
        result = cursor.fetchone()
        stats['total'] = result['total'] if result else 0
        
        cursor.execute("SELECT status, COUNT(*) as count FROM tickets GROUP BY status")
        status_results = cursor.fetchall()
        status_counts = {row['status']: row['count'] for row in status_results}
        
        stats['open'] = status_counts.get('Open', 0)
        stats['in_progress'] = status_counts.get('In Progress', 0)
        stats['resolved'] = status_counts.get('Resolved', 0)
        
        cursor.execute("""
            SELECT AVG(TIMESTAMPDIFF(HOUR, created_at, resolution_note_date)) as avg_hours 
            FROM tickets 
            WHERE status = 'Resolved' AND resolution_note_date IS NOT NULL
        """)
        avg_result = cursor.fetchone()
        avg_hours = avg_result['avg_hours'] if avg_result and avg_result['avg_hours'] else None
        stats['avg_resolution_time_hours'] = round(avg_hours, 2) if avg_hours else 0

        cursor.execute("""
            SELECT COUNT(*) as breaches 
            FROM tickets 
            WHERE status != 'Resolved' AND sla_deadline < NOW()
        """)
        breach_result = cursor.fetchone()
        stats['sla_breaches'] = breach_result['breaches'] if breach_result else 0
        
        cursor.execute("SELECT issue_type, COUNT(*) as count FROM tickets GROUP BY issue_type")
        issue_results = cursor.fetchall()
        stats['by_issue_type'] = {row['issue_type']: row['count'] for row in issue_results}

        cursor.execute("SELECT department, COUNT(*) as count FROM tickets GROUP BY department")
        dept_results = cursor.fetchall()
        stats['by_department'] = {row['department']: row['count'] for row in dept_results}
        
        print(f"Stats generated: {stats}")
        
        return jsonify(stats)
        
    except Error as e:
        print(f"Database error in ticket_stats: {e}")
        return jsonify({'error': 'Database error occurred'}), 500
    finally:
        if connection and connection.is_connected():
            connection.close()

@app.route('/api/hr_staff')
def get_hr_staff():
    connection = create_db_connection()
    if not connection:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = connection.cursor(dictionary=True)
        cursor.execute("SELECT name FROM hr_staff")
        staff = [row['name'] for row in cursor.fetchall()]
        return jsonify(staff)
    except Error as e:
        print(f"Database error: {e}")
        return jsonify({'error': 'Database error occurred'}), 500
    finally:
        if connection and connection.is_connected():
            connection.close()

@app.route('/api/tickets/clear', methods=['DELETE'])
def clear_all_tickets():
    connection = create_db_connection()
    if not connection:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = connection.cursor()
        cursor.execute("DELETE FROM tickets")
        connection.commit()
        return jsonify({'message': 'All tickets cleared successfully'})
    except Error as e:
        print(f"Database error: {e}")
        return jsonify({'error': 'Database error occurred'}), 500
    finally:
        if connection and connection.is_connected():
            connection.close()

@app.route('/api/reports/<report_type>', methods=['GET'])
def generate_report(report_type):
    connection = create_db_connection()
    if not connection:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = connection.cursor(dictionary=True)
        cursor.execute("""
            SELECT id, name, department, issue_type, priority, status, description, 
                   assigned_to, resolution_note, created_at, sla_deadline, resolution_note_date
            FROM tickets
            ORDER BY created_at DESC
        """)
        tickets = cursor.fetchall()

        if report_type == 'pdf':
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            c.setFont("Helvetica", 12)
            c.drawString(100, 750, "KEPROBA HR Ticket Report")
            y = 700
            for ticket in tickets:
                c.drawString(50, y, f"Ticket #{ticket['id']}: {ticket['issue_type']}")
                c.drawString(50, y-20, f"Submitted by: {ticket['name']} ({ticket['department']})")
                c.drawString(50, y-40, f"Status: {ticket['status']} | Priority: {ticket['priority']}")
                c.drawString(50, y-60, f"Created: {ticket['created_at'].strftime('%Y-%m-%d %H:%M:%S')}")
                y -= 80
                if y < 50:
                    c.showPage()
                    y = 750
            c.save()
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name='HR_Report.pdf', mimetype='application/pdf')

        elif report_type == 'docx':
            doc = Document()
            doc.add_heading('KEPROBA HR Ticket Report', 0)
            for ticket in tickets:
                doc.add_heading(f"Ticket #{ticket['id']}: {ticket['issue_type']}", level=1)
                doc.add_paragraph(f"Submitted by: {ticket['name']} ({ticket['department']})")
                doc.add_paragraph(f"Status: {ticket['status']} | Priority: {ticket['priority']}")
                doc.add_paragraph(f"Created: {ticket['created_at'].strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph(f"Description: {ticket['description']}")
                if ticket['resolution_note']:
                    doc.add_paragraph(f"Resolution: {ticket['resolution_note']}")
                doc.add_paragraph('')
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name='HR_Report.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        elif report_type == 'xlsx':
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = 'HR Tickets'
            headers = ['ID', 'Name', 'Department', 'Issue Type', 'Priority', 'Status', 'Description', 'Assigned To', 'Created At', 'SLA Deadline', 'Resolution Note']
            sheet.append(headers)
            for ticket in tickets:
                sheet.append([
                    ticket['id'],
                    ticket['name'],
                    ticket['department'],
                    ticket['issue_type'],
                    ticket['priority'],
                    ticket['status'],
                    ticket['description'],
                    ticket['assigned_to'] or 'Unassigned',
                    ticket['created_at'].strftime('%Y-%m-%d %H:%M:%S') if ticket['created_at'] else '',
                    ticket['sla_deadline'].strftime('%Y-%m-%d %H:%M:%S') if ticket['sla_deadline'] else '',
                    ticket['resolution_note'] or ''
                ])
            buffer = BytesIO()
            workbook.save(buffer)
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name='HR_Report.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

        else:
            return jsonify({'error': 'Invalid report type'}), 400

    except Error as e:
        print(f"Database error: {e}")
        return jsonify({'error': 'Database error occurred'}), 500
    finally:
        if connection and connection.is_connected():
            connection.close()

if __name__ == '__main__':
    #init_database()
    app.run(debug=True)